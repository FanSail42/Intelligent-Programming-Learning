# -*- coding: utf-8 -*-
"""Fill section 4 references into proposal DOCX with Word auto-numbering and renumber in-text citations."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "out_data" / "论文" / "01_本科毕业论文开题报告及开题答辩记录表.docx"
REF_MD = ROOT / "out_data" / "论文" / "开题报告撰写思路" / "07_主要参考文献.md"
SAMPLE_DOCX = ROOT / "out_data" / "论文" / "开题报告参考样例" / "03.docx"
BACKUP_PATH = DOCX_PATH.with_suffix(".docx.bak")

# Old in-text citation -> new sequential number (1..25)
OLD_TO_NEW: dict[str, int] = {
    "[1]": 1,
    "[2]": 2,
    "[3]": 3,
    "[4]": 4,
    "[5]": 5,
    "[6]": 6,
    "[7]": 7,
    "[8]": 8,
    "[10]": 9,
    "[13]": 10,
    "[15]": 11,
    "[16]": 12,
    "[17]": 13,
    "[18]": 14,
    "[19]": 15,
    "[22]": 16,
    "[23]": 17,
    "[24]": 18,
    "[25]": 19,
    "[27]": 20,
    "[E1]": 21,
    "[E2]": 22,
    "[E3]": 23,
    "[E4]": 24,
    "[E5]": 25,
}

REF_NUM_ID = "2"  # matches sample 03.docx bibliography list
REF_STYLE = "13"  # paragraph style used by sample reference entries


def load_reference_lines() -> list[str]:
    text = REF_MD.read_text(encoding="utf-8")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines[0].startswith("四、"):
        lines = lines[1:]
    return lines


def set_run_font(run, *, size_pt: float = 12, superscript: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    run.font.superscript = superscript
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def make_num_pr(num_id: str = REF_NUM_ID, ilvl: str = "0") -> OxmlElement:
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    return num_pr


def apply_reference_paragraph_format(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ppr.append(make_num_pr())
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "460")
    spacing.set(qn("w:lineRule"), "exact")
    ppr.append(spacing)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "0")
    ppr.append(ind)
    if REF_STYLE:
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), REF_STYLE)
        ppr.insert(0, pstyle)


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    cell._tc.append(OxmlElement("w:p"))


def add_title_paragraph(cell, title: str) -> None:
    p = cell.paragraphs[0]
    p.text = title
    for run in p.runs:
        set_run_font(run, size_pt=12, superscript=False)


def add_reference_entry(cell, text: str) -> None:
    p = cell.add_paragraph()
    apply_reference_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run, size_pt=12, superscript=False)


def replace_citations_in_cell(cell) -> int:
    """Replace [n]/[En] markers with superscript [new_n] runs."""
    count = 0
    pattern = re.compile(r"\[(?:E\d+|\d+)\]")

    for paragraph in cell.paragraphs:
        full_text = paragraph.text
        if not pattern.search(full_text):
            continue

        # Collect run formatting template
        template_rpr = None
        if paragraph.runs:
            template_rpr = deepcopy(paragraph.runs[0]._element.rPr)

        parts: list[tuple[str, bool]] = []
        pos = 0
        for m in pattern.finditer(full_text):
            if m.start() > pos:
                parts.append((full_text[pos : m.start()], False))
            old = m.group(0)
            new_n = OLD_TO_NEW.get(old)
            if new_n is None:
                parts.append((old, False))
            else:
                parts.append((f"[{new_n}]", True))
                count += 1
            pos = m.end()
        if pos < len(full_text):
            parts.append((full_text[pos:], False))

        # Rebuild paragraph
        p_el = paragraph._element
        for child in list(p_el):
            if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
                p_el.remove(child)

        for text, is_cite in parts:
            if not text:
                continue
            run = paragraph.add_run(text)
            if template_rpr is not None:
                run._element.insert(0, deepcopy(template_rpr))
            set_run_font(run, size_pt=12, superscript=is_cite)

    return count


def patch_bibliography_numbering() -> None:
    """Ensure numId=2 renders as [1][2]… (sample format), not Chinese enumeration."""
    import zipfile

    with zipfile.ZipFile(SAMPLE_DOCX) as zs:
        sample_num = zs.read("word/numbering.xml").decode("utf-8")
    with zipfile.ZipFile(DOCX_PATH) as zt:
        target_num = zt.read("word/numbering.xml").decode("utf-8")

    sample_abs0 = re.search(
        r'<w:abstractNum w:abstractNumId="0">.*?</w:abstractNum>', sample_num, re.S
    ).group(0)
    target_num = re.sub(
        r'<w:abstractNum w:abstractNumId="0">.*?</w:abstractNum>',
        sample_abs0,
        target_num,
        count=1,
        flags=re.S,
    )

    tmp = DOCX_PATH.with_suffix(".patch.docx")
    with zipfile.ZipFile(DOCX_PATH, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/numbering.xml":
                data = target_num.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(DOCX_PATH)


def main() -> None:
    refs = load_reference_lines()
    if len(refs) != 25:
        raise ValueError(f"Expected 25 references, got {len(refs)}")

    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    patch_bibliography_numbering()

    doc = Document(str(DOCX_PATH))
    table = doc.tables[0]

    # Renumber in-text citations in section 1 (row 1)
    replaced = replace_citations_in_cell(table.rows[1].cells[0])
    print(f"Renumbered {replaced} in-text citation markers in row 1")

    # Fill section 4 (row 4)
    ref_cell = table.rows[4].cells[0]
    clear_cell(ref_cell)
    add_title_paragraph(ref_cell, "四、主要参考文献")
    ref_cell.add_paragraph()  # blank line after title, like template
    for line in refs:
        add_reference_entry(ref_cell, line)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")
    print(f"Backup: {BACKUP_PATH}")
    print(f"Inserted {len(refs)} auto-numbered reference entries (numId={REF_NUM_ID})")


if __name__ == "__main__":
    main()
