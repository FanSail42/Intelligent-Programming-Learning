# -*- coding: utf-8 -*-
"""Fill thesis proposal DOCX template from markdown draft sections."""
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "root_data" / "02_本科毕业论文开题报告及开题答辩记录表.docx"
DST = ROOT / "out_data" / "20260611_本科毕业论文开题报告_初稿.docx"

TITLE = "基于大语言模型与检索增强生成的智能编程学习助教系统设计与实现"

SECTION1 = """（一）研究背景及意义

随着高等教育信息化与人工智能技术的快速发展，计算机类课程已成为高校人才培养体系中的核心组成部分。数据结构、面向对象程序设计、Python 数据分析等课程普遍采用“课堂讲授+编程实验+在线评测”的教学模式。然而，在实际教学过程中仍面临三方面突出矛盾：一是班级规模扩大导致教师难以及时、逐人解答编程疑问；二是通用搜索引擎与聊天机器人缺乏课程语境，学生获得的答案往往与教学大纲、课堂讲义脱节；三是传统在线评测系统（OJ）侧重“判对错”，对语法错误、逻辑缺陷与调试思路的讲解不足，难以支撑形成性评价与个性化复习。

近年来，以 GPT、通义千问、DeepSeek 等为代表的大语言模型（LLM）在代码理解、自然语言生成与多轮对话方面展现出较强能力；检索增强生成（RAG）通过将外部知识库与生成模型结合，可有效缓解大模型“幻觉”问题；学习分析则为刻画学生薄弱知识点、错题分布与复习路径提供了量化手段。

本课题拟设计并实现“慧编学伴”智能编程学习助教系统，面向高校计算机类课程场景，构建覆盖学生、教师、管理员三类角色的 Web 应用。系统以 FastAPI、Vue3、MySQL、Redis、Chroma 向量数据库及主流 LLM API 为技术基础，实现课程知识库管理、基于 RAG 的 AI 对话辅导、智能代码讲解、学习仪表盘与教师班级学情统计等功能。

理论意义：探索 LLM 与 RAG 在高校编程教学中的系统化集成方式，研究课程级知识隔离、引用溯源与相关性过滤等机制对辅导质量的影响；将规则驱动的学习分析与错题本联动相结合，为“AI 辅导—行为采集—学情反馈”闭环提供可复现的工程范式。

实践意义：为教师提供可管理的课程资料与班级学情视图，为学生提供 7×24 小时可及的智能辅导与学习档案，为管理员提供 AI 模型与用量管控能力，有助于缓解师资压力、提升实验课辅导效率。

（二）国内外研究现状

1. 智能编程教育与在线辅导系统：国外 Codecademy、Khan Academy 等平台强调即时反馈；国内多数系统仍以结构化题库为中心，对自然语言问答与个性化代码讲解支持有限。

2. 大语言模型在编程教学中的应用：LLM 可辅助代码解释与错误诊断，但也存在答案不可靠、与课程要求不一致等问题。教育领域逐渐转向“领域知识约束+可控提示工程+引用溯源”的专用助教架构。

3. 检索增强生成（RAG）在教育场景的研究：RAG 通过文档解析、分块、向量化与 Top-K 检索，将生成回答锚定在本地知识库上。研究重点正从“能否回答问题”转向“能否可追溯地引用教学材料”。

4. 学习分析与个性化推荐：代码提交、错误类型、AI 对话无上下文等事件可映射为知识点掌握度与错题本条目。基于规则引擎的方法具有可解释、可测试、成本低等优势。

5. 研究述评：现有研究与产品多在单点能力上较为成熟，但面向高校多课程、多角色、可运维的一体化智能编程助教系统仍相对缺乏。本课题强调课程 RAG 辅导、结构化代码讲解、规则化学情分析、教师班级视图与管理员 AI 治理的端到端闭环。"""

SECTION2 = """（一）研究内容

1. 需求分析与总体架构设计：面向学生、教师、管理员三类角色，梳理功能需求，设计前后端分离的 B/S 架构及 RESTful API 规范。

2. 课程知识库与 RAG 对话辅导模块：研究教学资料解析、分块、向量化与按课程隔离存储；实现语义检索、相关性过滤与 SSE 流式多轮对话，并展示引用来源。

3. 智能代码讲解模块：面向 C/C++、Java、Python 等语言，设计语法、语义、运行三级结构化讲解方案，并与 Monaco 编辑器集成。

4. 学习分析与推荐模块：研究学习事件埋点、错题本自动归集、知识点掌握度计算与规则推荐；实现学习仪表盘与错题本可视化。

5. 教师教学支持与管理模块：实现班级学情概览、账户管理、AI 模型配置、API 密钥加密存储与 Token 用量监控。

6. 系统测试与综合评估：建立 pytest 自动化测试体系，开展功能测试、接口测试与答辩演示验证。

（二）拟解决的关键问题

1. 课程语境下的可信 AI 辅导问题：如何避免通用大模型回答脱离课程讲义，并使学生能够核对引用来源。

2. 教学资料到可检索知识库的工程化流水线问题：如何实现资料上传、解析、向量化、失败重试与软删除，并保证按课程隔离向量数据。

3. 超越 OJ 的代码教学反馈问题：如何将 LLM 输出约束为结构化 JSON，形成可分级展示的讲解，并与错题本联动。

4. 可解释、可测试的学情建模问题：如何在不依赖 LLM 排序的前提下，基于学习事件计算掌握度并生成复习推荐。

5. 多角色权限与 AI 资源治理问题：如何在 JWT+RBAC 体系下隔离角色能力，并对 AI 调用进行限流、配额与密钥安全管理。

（三）拟取得的主要创新点

1. 面向课程的知识隔离式 RAG 辅导机制：整合向量检索、相关性过滤、引用展示与 SSE 流式输出，强化答案可追溯性。

2. “代码讲解—错题本—掌握度—推荐”规则闭环：以代码分析与对话行为为数据源，采用可解释的掌握度公式与 Top3 规则推荐。

3. 三级结构化智能代码讲解：将大模型能力约束为语法/语义/运行分层反馈，提升编程实验课辅导体验。

4. 兼顾教学与管理的多角色平台设计：打通学生学习端、教师班级学情端与管理员 AI/账户治理端。"""

SECTION3 = """（一）研究方法

1. 文献研究法：检索智能编程教育、LLM 教学应用、RAG、学习分析等相关文献与开源项目。

2. 需求调研法：结合高校计算机类课程教学流程，归纳学生、教师、管理员的核心用例。

3. 原型迭代法：采用 Phase 0～4 增量开发，优先完成学生端 P0 闭环，再扩展教师与管理功能。

4. 实验测试法：通过单元测试、接口测试、前后端联调与答辩脚本演示进行验证。

（二）步骤与技术路线

总体路线：需求分析→架构设计→数据库与 API 设计→后端服务实现→前端开发→联调测试→部署优化→论文撰写。

技术架构：前端 Vue3+Element Plus+ECharts+Monaco；后端 FastAPI+SQLAlchemy+Alembic+JWT；数据层 MySQL+Redis+Chroma；AI 层可配置 LLM API+Embedding+RAG 流水线。

图1 技术路线：文献调研→架构设计→后端模块（M01认证/M02课程/M03知识库/M04 RAG对话/M05代码讲解/M06学情/M07教师/M09管理）→前端多角色界面→pytest 自动化测试→部署与论文撰写。

（三）研究进度

1. 已完成工作量（截至2026年6月）：Phase 0～4 主体功能已完成；学生端核心闭环（登录、选课、RAG 对话、代码讲解、学习仪表盘、错题本、个人中心）可演示；教师班级学情、管理员 AI/账户管理已交付；后端 pytest 154 项全部通过；前端构建通过；模块说明与部署文档齐备。综合完成度约 84%。

2. 后续研究计划：2026年6月完善 M08 最小生产部署；7月系统测试与论文第3～5章撰写；8月完成论文第6～7章与答辩材料；9月定稿答辩。"""

SECTION4 = """[1] Brown T, et al. Language Models are Few-Shot Learners[C]//NeurIPS, 2020.
[2] Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//NeurIPS, 2020.
[3] Denny P, et al. Evaluating a LLM-Powered Programming Tutor[C]//ACE, 2023.
[4] Siemens G, Long P. Penetrating the Fog: Analytics in Learning and Education[J]. EDUCAUSE Review, 2011.
[5] 教育部. 教育信息化2.0行动计划[Z]. 2018.
[6] 张华, 李明. 基于知识图谱的智能导学系统研究综述[J]. 电化教育研究, 2022.
[7] 王磊, 陈静. 大语言模型在教育领域的应用现状与展望[J]. 开放教育研究, 2024.
[8] 刘洋, 赵强. 在线编程判题系统设计与实现[J]. 计算机应用与软件, 2021.
[9] Karpukhin V, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//EMNLP, 2020.
[10] FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2024.
[11] Vue.js Official Guide[EB/OL]. https://vuejs.org/, 2024.
[12] Chroma Documentation[EB/OL]. https://docs.trychroma.com/, 2024."""


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    table = doc.tables[0]

    table.rows[0].cells[1].text = TITLE
    set_cell_text(table.rows[1].cells[0], SECTION1)
    set_cell_text(table.rows[2].cells[0], SECTION2)
    set_cell_text(table.rows[3].cells[0], SECTION3)
    set_cell_text(table.rows[4].cells[0], SECTION4)

    doc.save(DST)
    body_len = len(SECTION1) + len(SECTION2) + len(SECTION3)
    print(f"Saved: {DST}")
    print(f"Body chars (sections 1-3): {body_len}")


if __name__ == "__main__":
    main()
