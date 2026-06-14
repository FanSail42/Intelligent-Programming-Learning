# -*- coding: utf-8 -*-
"""Renumber in-text citations [1].. by first appearance; reorder bibliography to match."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "out_data" / "论文" / "01_本科毕业论文开题报告及开题答辩记录表.docx"
REF_MD = ROOT / "out_data" / "论文" / "开题报告撰写思路" / "07_主要参考文献.md"
SAMPLE_DOCX = ROOT / "out_data" / "论文" / "开题报告参考样例" / "03.docx"
BACKUP_PATH = DOCX_PATH.with_suffix(".docx.bak2")

REF_NUM_ID = "2"
REF_STYLE = "13"
BODY_ROWS = (1, 2, 3)
REF_ROW = 4
CITE_RE = re.compile(r"\[\d+\]")


def load_reference_lines() -> list[str]:
    lines = [
        ln.strip()
        for ln in REF_MD.read_text(encoding="utf-8").splitlines()
        if ln.strip() and not ln.strip().startswith("四、")
    ]
    return lines


def set_run_font(run, *, size_pt: float = 12, superscript: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    run.font.superscript = superscript
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def make_num_pr(num_id: str = REF_NUM_ID, ilvl: str = "0") -> OxmlElement:
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    return num_pr


def apply_reference_paragraph_format(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), REF_STYLE)
    ppr.insert(0, pstyle)
    ppr.append(make_num_pr())
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "460")
    spacing.set(qn("w:lineRule"), "exact")
    ppr.append(spacing)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "0")
    ppr.append(ind)


def patch_bibliography_numbering() -> None:
    import zipfile

    with zipfile.ZipFile(SAMPLE_DOCX) as zs:
        sample_num = zs.read("word/numbering.xml").decode("utf-8")
    with zipfile.ZipFile(DOCX_PATH) as zt:
        target_num = zt.read("word/numbering.xml").decode("utf-8")

    sample_abs0 = re.search(
        r'<w:abstractNum w:abstractNumId="0">.*?</w:abstractNum>', sample_num, re.S
    ).group(0)
    target_num = re.sub(
        r'<w:abstractNum w:abstractNumId="0">.*?</w:abstractNum>',
        sample_abs0,
        target_num,
        count=1,
        flags=re.S,
    )

    tmp = DOCX_PATH.with_suffix(".patch.docx")
    with zipfile.ZipFile(DOCX_PATH, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/numbering.xml":
                data = target_num.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(DOCX_PATH)


def collect_first_appearance(table) -> list[int]:
    seen: set[int] = set()
    order: list[int] = []
    for ri in BODY_ROWS:
        text = table.rows[ri].cells[0].text
        for m in CITE_RE.finditer(text):
            n = int(m.group(0)[1:-1])
            if n not in seen:
                seen.add(n)
                order.append(n)
    return order


def build_reorder_maps(
    refs: list[str], appearance: list[int]
) -> tuple[dict[int, int], list[str]]:
    if len(appearance) != len(refs):
        missing = set(range(1, len(refs) + 1)) - set(appearance)
        extra = set(appearance) - set(range(1, len(refs) + 1))
        raise ValueError(f"Citation mismatch. missing={missing}, extra={extra}")

    old_to_new = {old: new for new, old in enumerate(appearance, start=1)}
    reordered = [refs[old - 1] for old in appearance]
    return old_to_new, reordered


def replace_citations_in_cell(cell, old_to_new: dict[int, int]) -> int:
    count = 0
    for paragraph in cell.paragraphs:
        full_text = paragraph.text
        if not CITE_RE.search(full_text):
            continue

        template_rpr = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs else None
        parts: list[tuple[str, bool]] = []
        pos = 0
        for m in CITE_RE.finditer(full_text):
            if m.start() > pos:
                parts.append((full_text[pos : m.start()], False))
            old_n = int(m.group(0)[1:-1])
            new_n = old_to_new.get(old_n)
            if new_n is None:
                parts.append((m.group(0), False))
            else:
                parts.append((f"[{new_n}]", True))
                count += 1
            pos = m.end()
        if pos < len(full_text):
            parts.append((full_text[pos:], False))

        p_el = paragraph._element
        for child in list(p_el):
            if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
                p_el.remove(child)

        for text, is_cite in parts:
            if not text:
                continue
            run = paragraph.add_run(text)
            if template_rpr is not None:
                run._element.insert(0, deepcopy(template_rpr))
            set_run_font(run, size_pt=12, superscript=is_cite)

    return count


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    cell._tc.append(OxmlElement("w:p"))


def write_bibliography(cell, refs: list[str]) -> None:
    clear_cell(cell)
    title_p = cell.paragraphs[0]
    title_p.text = "四、主要参考文献"
    for run in title_p.runs:
        set_run_font(run)
    cell.add_paragraph()
    for line in refs:
        p = cell.add_paragraph()
        apply_reference_paragraph_format(p)
        run = p.add_run(line)
        set_run_font(run)


def remove_body_list_numbering(table) -> None:
    """Remove bibliography numId=2 from body headings (would steal [1])."""
    for ri in BODY_ROWS:
        cell = table.rows[ri].cells[0]
        for para in cell.paragraphs:
            ppr = para._p.pPr
            if ppr is None or ppr.numPr is None:
                continue
            num_id = ppr.numPr.numId.val if ppr.numPr.numId is not None else None
            if str(num_id) == REF_NUM_ID:
                ppr.remove(ppr.numPr)


def main() -> None:
    refs = load_reference_lines()
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    patch_bibliography_numbering()

    doc = Document(str(DOCX_PATH))
    table = doc.tables[0]

    appearance = collect_first_appearance(table)
    old_to_new, reordered = build_reorder_maps(refs, appearance)

    remove_body_list_numbering(table)

    total = 0
    for ri in BODY_ROWS:
        total += replace_citations_in_cell(table.rows[ri].cells[0], old_to_new)

    write_bibliography(table.rows[REF_ROW].cells[0], reordered)
    doc.save(DOCX_PATH)

    print("First appearance order (old -> new):")
    for old, new in sorted(old_to_new.items(), key=lambda x: x[1]):
        print(f"  [{new}] <- was [{old}]")
    print(f"Updated {total} citation markers")
    print(f"First in-text citation is now [{appearance and old_to_new[appearance[0]]}]")
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
