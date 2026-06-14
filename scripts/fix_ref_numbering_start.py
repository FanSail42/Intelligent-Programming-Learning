# -*- coding: utf-8 -*-
"""Fix bibliography starting at [2]: remove stray numId=2 from section headings."""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

DOCX_PATH = Path(__file__).resolve().parents[1] / "out_data" / "论文" / "01_本科毕业论文开题报告及开题答辩记录表.docx"
BODY_ROWS = (0, 1, 2, 3)
REF_ROW = 4


def remove_num_pr_if_bibliography_list(paragraph) -> bool:
    """Only strip numId=2 from body; leave other list styles (1. 2. 3.) intact."""
    ppr = paragraph._p.pPr
    if ppr is None or ppr.numPr is None:
        return False
    num_id = ppr.numPr.numId.val if ppr.numPr.numId is not None else None
    if str(num_id) != "2":
        return False
    ppr.remove(ppr.numPr)
    return True


def main() -> None:
    doc = Document(str(DOCX_PATH))
    table = doc.tables[0]
    removed = []

    for ri in BODY_ROWS:
        cell = table.rows[ri].cells[0]
        for pi, para in enumerate(cell.paragraphs):
            if remove_num_pr_if_bibliography_list(para):
                removed.append(f"row{ri}/p{pi}: {para.text[:40]!r}")

    doc.save(DOCX_PATH)
    print(f"Removed numPr from {len(removed)} paragraph(s):")
    for item in removed:
        print(" ", item)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
